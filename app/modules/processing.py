import os
from  faster_whisper import WhisperModel, BatchedInferencePipeline
import re
from app.modules.gemini_ai import refine_transcript
from docx import Document



def clean_text(text: str) -> str:
    """
    Thực hiện tiền xử lý văn bản:
      - Loại bỏ khoảng trắng thừa giữa các từ.
      - Cắt bỏ khoảng trắng đầu/cuối.
    Có thể mở rộng thêm các bước xử lý (ví dụ: chuẩn hóa dấu câu) nếu cần.
    """
    text = re.sub(r'\s+', ' ', text)
    return text.strip()





def preprocess_transcript(segments: list):
    """
    Tiền xử lý transcript, trả về dict với data:
    -text: đoạn văn được transcript
    """

    processed_segments = []
    for segment in segments:
        processed_segments.append({
            'start': segment.start,
            'end': segment.end,
            'text': segment.text,
        })

    return processed_segments


def transcript_audio (
        input_audio: str = "audio.mp3",
        model_size: str = "base",
        device: str = "cpu",
        compute_type: str = "int8",
        beam_size: int = 8,
        vad_filter: bool = False,
):
    """
    
    :param input_audio: 
    :param model_size: 
    :param device: 
    :param compute_type:  float 16,32 (chứa nhiều thông tin -> chính xác hơn)
    :param beam_size:  tăng độ chính xác
    :param vad_filter: nhận diện xem có giọng nói trong file không
    :return: 
    """

    if not os.path.exists(input_audio):
        raise FileNotFoundError(f"File '{input_audio}' không tồn tại.")

        # Khởi tạo model Faster Whisper với các tham số lấy từ config
    model = WhisperModel(model_size, device=device, compute_type=compute_type)

    # Cấu hình tham số cho quá trình transcription
    transcription_kwargs = {"beam_size": beam_size}
    if vad_filter:
        transcription_kwargs["vad_filter"] = True

    # Chạy quá trình transcription
    batched_model = BatchedInferencePipeline(model=model)
    segments, info = batched_model.transcribe(input_audio, **transcription_kwargs, batch_size=32)
    segments = list(segments)  # Ép generator thành list để dễ xử lý lại sau này
    processed_segments = preprocess_transcript(segments)
    return processed_segments, info




def save_transcript(segments: list, output_file: str, api_key: str) -> None:
    from math import ceil

    doc = Document()
    batch_size = 50  # Số câu mỗi lần gọi API

    # Chia thành từng batch 50 câu
    total_segments = len(segments)
    num_batches = ceil(total_segments / batch_size)

    for batch_index in range(num_batches):
        batch = segments[batch_index * batch_size : (batch_index + 1) * batch_size]

        # Tạo văn bản gộp để gửi lên Gemini
        numbered_lines = [f"{i+1}. {seg['text'].strip()}" for i, seg in enumerate(batch)]
        combined_text = "\n".join(numbered_lines)

        if api_key:
            try:
                refined_text = refine_transcript(combined_text, api_key)
                refined_lines = refined_text.splitlines()

                # Kiểm tra nếu số dòng trả về khác số segment gốc
                if len(refined_lines) != len(batch):
                    raise ValueError(f"Số dòng trả về ({len(refined_lines)}) không khớp số câu gốc ({len(batch)}).")

            except Exception as e:
                refined_lines = [seg["text"] for seg in batch]  # fallback nếu lỗi API

        else:
            refined_lines = [seg["text"] for seg in batch]

        # Ghi từng câu vào file .docx
        for i, seg in enumerate(batch):
            time_range = f"[{seg['start']:.2f}s → {seg['end']:.2f}s]"
            paragraph = doc.add_paragraph(time_range + " ")
            paragraph.add_run(refined_lines[i])

    doc.save(output_file)